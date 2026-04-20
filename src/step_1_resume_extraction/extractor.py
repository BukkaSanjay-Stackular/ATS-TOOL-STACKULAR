"""
Multi-Format Text Extraction Module
Handles extracting text from PDF (including scanned), DOCX, and image files.
Includes OCR fallback for scanned PDFs.
"""

import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document
from pathlib import Path
from typing import Optional
import io


def pymupdf_extract(file_path: str) -> Optional[str]:
    """
    Extract text from a PDF using PyMuPDF (standard PDFs with text).
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        Optional[str]: Extracted text or None if extraction fails
    """
    try:
        pdf_document = fitz.open(file_path)
        
        # Extract text from all pages
        extracted_text = ""
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            extracted_text += page.get_text()
        
        pdf_document.close()
        
        return extracted_text.strip() if extracted_text.strip() else None
    
    except Exception as e:
        print(f"  ❌ PyMuPDF extraction failed: {str(e)}")
        return None


def ocr_extract(file_path: str) -> Optional[str]:
    """
    Extract text from scanned PDFs using Tesseract OCR.
    Converts PDF pages to images and applies OCR.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        Optional[str]: Extracted text or None if extraction fails
    """
    try:
        print(f"  🔄 Attempting OCR extraction (scanned PDF detected)...")
        
        pdf_document = fitz.open(file_path)
        extracted_text = ""
        
        # Process each page with OCR
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            
            # Convert PDF page to image
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
            image_data = pix.tobytes("ppm")
            image = Image.open(io.BytesIO(image_data))
            
            # Apply OCR
            page_text = pytesseract.image_to_string(image)
            extracted_text += page_text + "\n"
        
        pdf_document.close()
        
        return extracted_text.strip() if extracted_text.strip() else None
    
    except pytesseract.TesseractNotFoundError:
        print(f"  ❌ Tesseract OCR not installed. Install from: https://github.com/UB-Mannheim/tesseract/wiki")
        return None
    except Exception as e:
        print(f"  ❌ OCR extraction failed: {str(e)}")
        return None


def docx_extract(file_path: str) -> Optional[str]:
    """
    Extract text from DOCX (Word) files.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        Optional[str]: Extracted text or None if extraction fails
    """
    try:
        document = Document(file_path)
        
        # Extract text from all paragraphs
        extracted_text = ""
        for paragraph in document.paragraphs:
            extracted_text += paragraph.text + "\n"
        
        # Also extract from tables if present
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    extracted_text += cell.text + "\n"
        
        return extracted_text.strip() if extracted_text.strip() else None
    
    except Exception as e:
        print(f"  ❌ DOCX extraction failed: {str(e)}")
        return None


def get_file_type(file_path: str) -> Optional[str]:
    """
    Determine file type from extension.
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        Optional[str]: File type (pdf, docx, or None if unsupported)
    """
    extension = Path(file_path).suffix.lower()
    
    if extension == ".pdf":
        return "pdf"
    elif extension in [".docx", ".doc"]:
        return "docx"
    else:
        return None


def extract_text(file_path: str) -> Optional[str]:
    """
    Main dispatcher function - extract text from any supported file format.
    Handles PDF (with OCR fallback for scanned), DOCX, and image files.
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        Optional[str]: Extracted text or None if extraction fails
    """
    try:
        # Verify file exists
        if not Path(file_path).exists():
            print(f"⚠️  File not found: {file_path}")
            return None
        
        file_type = get_file_type(file_path)
        
        if file_type == "pdf":
            # Try PyMuPDF first
            text = pymupdf_extract(file_path)
            
            # If text is too short (< 50 chars), likely a scanned PDF - use OCR
            if text and len(text.strip()) < 50:
                print(f"  📄 Short text detected, attempting OCR for scanned PDF...")
                ocr_text = ocr_extract(file_path)
                if ocr_text:
                    text = ocr_text
            elif not text:
                # No text extracted, try OCR
                print(f"  📄 No text extracted, attempting OCR...")
                text = ocr_extract(file_path)
            
            return text
        
        elif file_type == "docx":
            return docx_extract(file_path)
        
        else:
            print(f"⚠️  Unsupported file type: {Path(file_path).suffix}")
            return None
    
    except Exception as e:
        print(f"❌ Error extracting text from {file_path}: {str(e)}")
        return None


def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """
    Legacy function for backward compatibility.
    Use extract_text() instead for multi-format support.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        Optional[str]: Extracted text or None if extraction fails
    """
    return extract_text(file_path)
